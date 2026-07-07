"""Best-effort text extraction for file attachments (MVP context-stuffing, not RAG).

Extracted text is injected verbatim into the LLM prompt by app.services.llm — see that
module's docstring and docs/ARCHITECTURE.md "LLM integration" for how it's used.
"""

import io
import logging

import pdfplumber
from docx import Document
from openpyxl import load_workbook

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "gif", "webp", "bmp"}
PDF_EXTENSIONS = {"pdf"}
DOCX_EXTENSIONS = {"docx"}
XLSX_EXTENSIONS = {"xlsx", "xlsm"}

_DOCX_MIMETYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_XLSX_MIMETYPES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
}


def classify_kind(filename: str, mimetype: str | None) -> str:
    """Returns one of "image" | "pdf" | "docx" | "xlsx" | "other", used to decide both storage
    metadata and which extractor (if any) to run.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    mimetype = mimetype or ""

    if ext in IMAGE_EXTENSIONS or mimetype.startswith("image/"):
        return "image"
    if ext in PDF_EXTENSIONS or mimetype == "application/pdf":
        return "pdf"
    if ext in DOCX_EXTENSIONS or mimetype in _DOCX_MIMETYPES:
        return "docx"
    if ext in XLSX_EXTENSIONS or mimetype in _XLSX_MIMETYPES:
        return "xlsx"
    return "other"


def _extract_pdf_text(data: bytes) -> str:
    parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
    return "\n\n".join(parts)


def _extract_docx_text(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    lines = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(lines)


def _extract_xlsx_text(data: bytes) -> str:
    workbook = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    lines = []
    for sheet in workbook.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                lines.append(", ".join("" if cell is None else str(cell) for cell in row))
    return "\n".join(lines)


_EXTRACTORS = {
    "pdf": _extract_pdf_text,
    "docx": _extract_docx_text,
    "xlsx": _extract_xlsx_text,
}


def extract_text(data: bytes, kind: str) -> str | None:
    """Returns extracted text for pdf/docx/xlsx, or None for kinds with nothing to extract
    (images, other). Extraction is best-effort: a malformed file logs and returns None rather
    than failing the upload.
    """
    extractor = _EXTRACTORS.get(kind)
    if extractor is None:
        return None
    try:
        return extractor(data)
    except Exception:
        logger.exception("Text extraction failed for kind=%s", kind)
        return None
